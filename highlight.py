import docx
from docx.enum.text import WD_COLOR_INDEX
import os
def highlighted(documents,subject,title):
    doc_endsem = docx.Document()
    doc_placement = docx.Document()
    for document in documents:
        if os.path.exists(document) and os.path.isfile(document):
            temp = docx.Document(document)
            for p in temp.paragraphs:
                i = 0
                j = 0
                k = 0
                for run in p.runs:
                    font = run.font
                    if font.highlight_color == WD_COLOR_INDEX.YELLOW:
                        if i == 0:
                            para1 = doc_endsem.add_paragraph()
                        if run.text.isprintable():
                            para1.add_run(run.text)
                        else:
                            run1 = para1.add_run(run.text)
                            font1 = run1.font
                            font1.name = 'Symbol'
                        i += 1
                    elif font.highlight_color == WD_COLOR_INDEX.RED:
                        if j == 0:
                            para2 = doc_placement.add_paragraph()
                        if run.text.isprintable():
                            para1.add_run(run.text)
                        else:
                            run1 = para1.add_run(run.text)
                            font1 = run1.font
                            font1.name = 'Symbol'
                        j += 1
                    elif font.highlight_color == WD_COLOR_INDEX.BLUE:
                        if k == 0:
                            para1 = doc_endsem.add_paragraph()
                            para2 = doc_placement.add_paragraph()
                        if run.text.isprintable():
                            para1.add_run(run.text)
                            para2.add_run(run.text)
                        else:
                            run1 = para1.add_run(run.text)
                            font1 = run1.font
                            font1.name = 'Symbol'
                            run1 = para2.add_run(run.text)
                            font1 = run1.font
                            font1.name = 'Symbol'
                        k += 1
    path = './Upload Folder/' + subject + '/' + 'Event Based_' + subject + '/'
    doc_endsem.save(path + title + '_ENDSEM.docx')
    doc_placement.save(path + title + '_PLACEMENT.docx')  

if __name__ == '__main__':
    path1 = os.path.join('Upload Folder','DBMS','temp')
    files = [os.path.join('Upload Folder','DBMS','temp',f) for f in os.listdir(path1)]
    highlighted(files,'DBMS','test1')
    for f in files:
        os.remove(f)


   